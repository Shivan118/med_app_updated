import streamlit as st
import google.generativeai as genai
from google_api_key import google_api_key
from PIL import Image
import io
import docx
from io import BytesIO
from analysis.ultrasound_analyzer import configure_ultrasound_model, analyze_ultrasound, generate_ultrasound_doc
from analysis.xrays_analyzer import configure_xrays_model, analyze_xrays, generate_xrays_doc
from analysis.ecg_analyzer import configure_ecg_model, analyze_ecg, generate_ecg_doc

# Configure Google Generative AI
genai.configure(api_key=google_api_key)

ecg_model = configure_ecg_model(google_api_key)
xrays_model = configure_xrays_model(google_api_key)
ultrasound_model = configure_ultrasound_model(google_api_key)

generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 0,
    "max_output_tokens": 8192,
}

safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
]

system_prompts = {
    "image": """
    You are a domain expert in medical image analysis. Examine the medical image and provide:
    1. Detailed Analysis: Scrutinize the image for abnormalities.
    2. Analysis Report: Document findings in a structured format.
    3. Recommendations: Suggest remedies, tests, or treatments.
    4. Treatments: If applicable, detail treatments for faster recovery.
    
    Important Notes:
    1. Only respond to human health-related images.
    2. Note if aspects are 'Unable to be correctly determined based on the uploaded image'
    3. Include the disclaimer: "Caution: This is an AI BOT made by MED360. Consult with a Doctor before making any decisions."
    
    Provide the response in headings and sub-headings in bullet format.
    Extend output to the next line when encountering \n or \n\n in responses.
    """,
    "text": """
    You are an AI medical assistant. Provide concise, accurate information on medical topics. Structure your answer with:
    1. Brief Overview
    2. Key Symptoms (3-5)
    3. Medication Suggestions (2-3 common medications, purpose, dosage forms)
    4. Quick Suggestions (2-3 health tips)
    5. When to Seek Help
    
    Keep responses within 200-250 words. Use simple language. Preface medication suggestions with: "Common medications that a doctor might consider include: "
    End with: "Note: This is AI-generated information by MED360. Always consult a healthcare professional for personalized medical advice."
    """
}

model = genai.GenerativeModel(model_name="gemini-1.5-pro-latest",
                              generation_config=generation_config,
                              safety_settings=safety_settings)

def classify_medical_image(image, model):
    image_data = io.BytesIO()
    image.save(image_data, format='JPEG')
    image_data = image_data.getvalue()
    image_parts = [{"mime_type": "image/jpeg", "data": image_data}]
    
    classification_prompt = """
    Classify the given image into one of these categories:
    1. Ultrasound
    2. X-ray
    3. ECG
    4. MRI
    5. CT Scan
    6. Mammogram
    7. Dental X-ray
    8. Retinal Scan
    9. Microscopic Image
    10. Other Medical Image

    Respond with ONLY the category name. If unsure, respond with "Unclassified".
    """
    
    prompt_parts = [image_parts[0], classification_prompt]
    response = model.generate_content(prompt_parts)
    return response.text.strip()

def analyze_image(image, image_type):
    if image_type == "ECG":
        return analyze_ecg(image, ecg_model), generate_ecg_doc
    elif image_type == "X-ray":
        return analyze_xrays(image, xrays_model), generate_xrays_doc
    elif image_type == "Ultrasound":
        return analyze_ultrasound(image, ultrasound_model), generate_ultrasound_doc
    else:
        # Fallback to generic analysis
        return system_prompts(image, model), generate_doc

def generate_doc(content):
    doc = docx.Document()
    sections = content.split('**')
    for i, section in enumerate(sections):
        if i % 2 != 0:
            p = doc.add_paragraph()
            p.add_run(section.strip()).bold = True
            if i + 1 < len(sections):
                content = sections[i + 1]
                for point in content.split('*'):
                    if point.strip():
                        doc.add_paragraph(point.strip(), style='List Bullet')
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

st.title("Visual Medical Assistant 👨‍⚕️ 🩺 🏥")
st.subheader("An app to help with medical analysis using images and text")

tab1, tab2 = st.tabs(["Image Analysis", "Text Query"])

with tab1:
    st.header("Image Analysis")
    uploaded_file = st.file_uploader("Upload the image for Analysis:", type=["jpg", "jpeg", "png"])

    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption="Uploaded Image", use_column_width=True)
        if st.button("Analyze"):
            try:
                with st.spinner("Classifying image..."):
                    image_type = classify_medical_image(image, model)
                    st.write(f"Image classified as: {image_type}")
                
                with st.spinner("Analyzing..."):
                    ai_response, doc_generator = analyze_image(image, image_type)
                    
                    if ai_response:
                        st.markdown(ai_response)
                        
                        doc_io = doc_generator(ai_response)
                        st.download_button(
                            label="Download analysis as .doc",
                            data=doc_io,
                            file_name=f"medical_image_analysis_{image_type}.doc",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.error("No response received from the analysis.")
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

with tab2:
    st.header("Text Query")
    query = st.text_area("Enter your medical query:")
    if st.button("Submit Query"):
        with st.spinner("Processing..."):
            prompt_parts = [
                system_prompts["text"],
                f"User query: {query}"
            ]
            ai_response = model.generate_content(prompt_parts)
            if ai_response:
                st.markdown(ai_response.text)
                
                doc_io = generate_doc(ai_response.text)
                st.download_button(
                    label="Download response as .doc",
                    data=doc_io,
                    file_name="medical_query_response.doc",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("An error occurred while processing your query.")

st.sidebar.title("About")
st.sidebar.info("This is a Visual Medical Assistant app that uses AI to analyze medical images and answer medical queries.")