import google.generativeai as genai
from PIL import Image
import io
import docx
from io import BytesIO

def configure_ultrasound_model(api_key):
    genai.configure(api_key=api_key)
    
    generation_config = {
        "temperature": 1,
        "top_p": 0.95,
        "top_k": 0,
        "max_output_tokens": 8192,
    }

    safety_settings = [
        {
            "category": "HARM_CATEGORY_HARASSMENT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_HATE_SPEECH",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
        {
            "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
            "threshold": "BLOCK_MEDIUM_AND_ABOVE"
        },
    ]

    return genai.GenerativeModel(model_name="gemini-1.5-pro-latest",
                                 generation_config=generation_config,
                                 safety_settings=safety_settings)

ultrasound_prompt = """
You are the world's top specialized ultrasound analyst and performed 1000 ultrasound across the world. You are known for giving the best report after analyzing the ultrasound. 
You are a specialized ultrasound analyst. Your task is to examine ultrasound images of the abdominal region and provide a detailed report. Focus on the following organs and structures, providing specific information for each:

Add a Heading "Ultrasound ANALYSIS REPORT" in Bold Letters, centered at the top of the first page.

PATIENT INFORMATION:
   - Patient Name:  
   - Age: 
   - Date: 
   - Time: 

1. LIVER:
   - Size, contour, and echotexture
   - Presence or absence of focal lesions
   - Condition of intrahepatic venous radicles

2. PORTAL VEIN:
   - Course and caliber

3. GALL BLADDER:
   - Size, contour, and wall thickness
   - Presence or absence of calculi

4. PANCREAS:
   - Size, contour, and echotexture
   - Presence or absence of focal lesions

5. SPLEEN:
   - Size, contour, and echotexture
   - Presence or absence of focal lesions

6. AORTA & IVC:
   - General condition

7. KIDNEYS (BOTH):
   - Size, contour, position, and echogenicity
   - Cortical thickness
   - Presence or absence of hydronephrosis or calculi
   - Condition of perirenal planes
   - Measurements of right and left kidneys

8. URINARY BLADDER:
   - Contour, capacity, and wall thickness
   - Presence or absence of calculi

9. PROSTATE (if visible):
   - Volume in cc
   - Size and echotexture
   - Presence or absence of focal lesions

10. ASCITES:
    - Presence or absence

For each structure, provide a detailed description using the following format:

**[ORGAN/STRUCTURE NAME]:** [Detailed description based on the points above]

If any measurement or specific detail is not visible or cannot be determined from the image, simply say ..... in your remport.

At the end of your report, provide an "IMPRESSION" section summarizing any significant findings or stating if the study appears within normal limits.

Remember to maintain a professional and objective tone. End your report with the following disclaimer:
"Caution: This analysis is generated by an AI system (MED360). For accurate diagnosis and treatment, please consult with a qualified healthcare professional."
"""

def analyze_ultrasound(image, model):
    image_data = io.BytesIO()
    image.save(image_data, format='JPEG')
    image_data = image_data.getvalue()
    image_parts = [{"mime_type": "image/jpeg", "data": image_data}]
    prompt_parts = [image_parts[0], ultrasound_prompt]
    response = model.generate_content(prompt_parts)
    return response.text

def generate_ultrasound_doc(content):
    doc = docx.Document()
    
    # Split the content into sections
    sections = content.split('**')
    
    for i, section in enumerate(sections):
        if i % 2 == 0:  # Even indexes are normal text or empty
            continue
        else:  # Odd indexes are headers or content
            # Add the header in bold
            p = doc.add_paragraph()
            p.add_run(section.strip()).bold = True
            
            # If there's content following this header, add it as bullet points
            if i + 1 < len(sections):
                content = sections[i + 1]
                bullet_points = content.split('*')
                for point in bullet_points:
                    if point.strip():
                        doc.add_paragraph(point.strip(), style='List Bullet')
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io